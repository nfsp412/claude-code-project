#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档转Markdown转换脚本

将.docx格式的Word文档转换为Markdown格式(.md)
支持包含altChunk嵌入HTML内容的文档
"""

import argparse
import os
import sys
import re
import zipfile
import quopri
import html
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from bs4 import BeautifulSoup


class WordToMarkdownConverter:
    """Word文档转Markdown转换器"""

    def __init__(self, docx_path):
        """
        初始化转换器

        Args:
            docx_path: Word文档路径
        """
        self.docx_path = Path(docx_path)
        self.document = None
        self.markdown_lines = []
        self.image_counter = 1
        self.has_alt_chunk = False

    def _has_alt_chunk(self):
        """检查文档是否包含altChunk（嵌入的HTML）"""
        try:
            with zipfile.ZipFile(str(self.docx_path), 'r') as z:
                # 检查document.xml中是否有altChunk
                if 'word/document.xml' in z.namelist():
                    with z.open('word/document.xml') as f:
                        content = f.read().decode('utf-8')
                        return '<w:altChunk' in content
        except Exception:
            pass
        return False

    def _extract_html_from_alt_chunk(self):
        """
        从altChunk中提取HTML内容并转换为Markdown

        Returns:
            转换后的Markdown内容
        """
        try:
            with zipfile.ZipFile(str(self.docx_path), 'r') as z:
                if 'word/afchunk.mht' in z.namelist():
                    with z.open('word/afchunk.mht') as f:
                        mht_content = f.read()

                    # 解码quoted-printable内容
                    decoded = quopri.decodestring(mht_content).decode('utf-8', errors='ignore')

                    # 提取HTML部分（在------=mhtDocumentPart分隔符之后）
                    html_match = re.search(r'------=mhtDocumentPart.*?<!DOCTYPE html>(.*?)<html>', decoded, re.DOTALL | re.IGNORECASE)
                    if html_match:
                        # 提取整个HTML文档
                        html_start = decoded.find('<!DOCTYPE html>')
                        if html_start == -1:
                            html_start = decoded.find('<html>')
                        if html_start != -1:
                            html_content = decoded[html_start:]
                            return self._html_to_markdown(html_content)

                    # 如果没有DOCTYPE，尝试提取body内容
                    body_match = re.search(r'<body[^>]*>(.*?)</body>', decoded, re.DOTALL | re.IGNORECASE)
                    if body_match:
                        return self._html_to_markdown(body_match.group(1))

                    # 最后尝试直接转换所有内容
                    return self._html_to_markdown(decoded)

        except Exception as e:
            print(f"警告: 无法提取altChunk内容: {e}")

        return ""

    def _html_to_markdown(self, html_content):
        """
        将HTML内容转换为Markdown

        Args:
            html_content: HTML字符串

        Returns:
            Markdown字符串
        """
        # 使用BeautifulSoup解析HTML
        soup = BeautifulSoup(html_content, 'html.parser')

        # 移除style和script标签
        for tag in soup(['style', 'script', 'meta', 'head']):
            tag.decompose()

        # 获取body内容，如果没有body则使用整个内容
        body = soup.find('body')
        content = body if body else soup

        result = []
        list_stack = []  # 用于跟踪嵌套列表

        def process_element(element, depth=0):
            """递归处理HTML元素"""
            if element.name is None:  # 文本节点
                text = str(element).strip()
                if text:
                    return text
                return ""

            # 处理标题
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                text = element.get_text().strip()
                return f"\n{'#' * level} {text}\n"

            # 处理段落
            elif element.name == 'p':
                text = ''.join(process_element(child, depth) for child in element.children).strip()
                if text:
                    return f"\n{text}\n"
                return ""

            # 处理粗体
            elif element.name == 'strong' or element.name == 'b':
                text = ''.join(process_element(child, depth) for child in element.children)
                return f"**{text}**"

            # 处理斜体
            elif element.name == 'em' or element.name == 'i':
                text = ''.join(process_element(child, depth) for child in element.children)
                return f"*{text}*"

            # 处理代码
            elif element.name == 'code':
                text = element.get_text().strip()
                return f"`{text}`"

            # 处理预格式化文本
            elif element.name == 'pre':
                text = element.get_text()
                return f"\n```\n{text}\n```\n"

            # 处理有序列表
            elif element.name == 'ol':
                items = []
                for i, li in enumerate(element.find_all('li', recursive=False), 1):
                    li_text = ' '.join(process_element(child, depth + 1) for child in li.children).strip()
                    items.append(f"{i}. {li_text}")
                return "\n" + "\n".join(items) + "\n"

            # 处理无序列表
            elif element.name == 'ul':
                items = []
                for li in element.find_all('li', recursive=False):
                    li_text = ' '.join(process_element(child, depth + 1) for child in li.children).strip()
                    items.append(f"- {li_text}")
                return "\n" + "\n".join(items) + "\n"

            # 处理列表项
            elif element.name == 'li':
                text = ''.join(process_element(child, depth) for child in element.children).strip()
                return text

            # 处理换行
            elif element.name == 'br':
                return "\n"

            # 处理链接
            elif element.name == 'a':
                text = element.get_text().strip()
                href = element.get('href', '')
                if href:
                    return f"[{text}]({href})"
                return text

            # 处理表格
            elif element.name == 'table':
                rows = []
                thead = element.find('thead')
                if thead:
                    header_cells = thead.find_all(['th', 'td'])
                    if header_cells:
                        header = "| " + " | ".join(cell.get_text().strip() for cell in header_cells) + " |"
                        separator = "| " + " | ".join("---" for _ in header_cells) + " |"
                        rows.append(header)
                        rows.append(separator)

                tbody = element.find('tbody') or element
                for tr in tbody.find_all('tr', recursive=False):
                    cells = tr.find_all(['td', 'th'], recursive=False)
                    if cells:
                        row = "| " + " | ".join(cell.get_text().strip() for cell in cells) + " |"
                        rows.append(row)

                return "\n" + "\n".join(rows) + "\n"

            # 处理div和span等容器
            elif element.name in ['div', 'span', 'section', 'article']:
                text = ''.join(process_element(child, depth) for child in element.children).strip()
                if text:
                    return f"\n{text}\n"
                return ""

            # 其他元素，递归处理子元素
            else:
                return ''.join(process_element(child, depth) for child in element.children)

        # 处理整个内容
        markdown = process_element(content)

        # 清理：移除多余的空行
        lines = markdown.split('\n')
        cleaned_lines = []
        prev_empty = False
        for line in lines:
            if line.strip():
                cleaned_lines.append(line)
                prev_empty = False
            elif not prev_empty:
                cleaned_lines.append('')
                prev_empty = True

        return '\n'.join(cleaned_lines).strip()

    def load_document(self):
        """加载Word文档"""
        if not self.docx_path.exists():
            raise FileNotFoundError(f"文件不存在: {self.docx_path}")

        # 检查是否有altChunk
        self.has_alt_chunk = self._has_alt_chunk()

        if self.has_alt_chunk:
            print(f"检测到文档包含嵌入的HTML内容（altChunk）")
            markdown_content = self._extract_html_from_alt_chunk()
            if markdown_content:
                self.markdown_lines = markdown_content.split('\n')
                print(f"成功从altChunk提取内容")
            else:
                print(f"警告: 无法从altChunk提取内容")
            # 仍然加载文档对象（虽然可能为空）
            try:
                self.document = Document(str(self.docx_path))
            except Exception:
                self.document = Document()
        else:
            try:
                self.document = Document(str(self.docx_path))
                print(f"成功加载文档: {self.docx_path.name}")
            except Exception as e:
                raise ValueError(f"无法加载Word文档: {e}")

    def iter_block_items(self):
        """
        遍历文档中的所有块级元素（段落和表格）

        Yields:
            段落或表格对象
        """
        if self.document is None:
            raise ValueError("文档未加载，请先调用 load_document()")

        parent = self.document.element.body
        for element in parent.iterchildren():
            if isinstance(element, CT_P):
                yield Paragraph(element, self.document)
            elif isinstance(element, CT_Tbl):
                yield Table(element, self.document)

    def _is_heading(self, paragraph):
        """判断段落是否为标题"""
        if paragraph.style is None:
            return False
        return paragraph.style.type == WD_STYLE_TYPE.PARAGRAPH and 'Heading' in paragraph.style.name

    def _get_heading_level(self, paragraph):
        """获取标题级别（1-6）"""
        style_name = paragraph.style.name
        level = int(style_name.replace('Heading ', '')) if 'Heading ' in style_name else 1
        return max(1, min(6, level))

    def _process_paragraph(self, paragraph):
        """
        处理段落，转换为Markdown格式

        Args:
            paragraph: docx.Paragraph对象
        """
        # 跳过空段落
        if not paragraph.text.strip():
            return

        # 检查是否为标题
        if self._is_heading(paragraph):
            level = self._get_heading_level(paragraph)
            text = self._clean_text(paragraph.text)
            self.markdown_lines.append(f"{'#' * level} {text}")
            return

        # 处理普通段落，保留格式
        markdown_text = self._convert_paragraph_to_markdown(paragraph)
        self.markdown_lines.append(markdown_text)

    def _clean_text(self, text):
        """清理文本，去除多余空格"""
        return ' '.join(text.split())

    def _convert_paragraph_to_markdown(self, paragraph):
        """
        将段落转换为Markdown格式，保留粗体、斜体等格式

        Args:
            paragraph: docx.Paragraph对象

        Returns:
            Markdown格式的文本
        """
        text = ""
        for run in paragraph.runs:
            run_text = run.text
            if not run_text:
                continue

            # 添加格式标记
            prefix = []
            suffix = []

            if run.bold:
                prefix.append("**")
                suffix.append("**")
            if run.italic:
                prefix.append("*")
                suffix.append("*")

            text += ''.join(prefix) + run_text + ''.join(reversed(suffix))

        # 检查是否为列表项
        if paragraph.style and 'List' in paragraph.style.name:
            # 判断是有序还是无序列表
            if 'Number' in paragraph.style.name or 'Decimal' in paragraph.style.name:
                # 有序列表 - 简化处理，使用1.
                return f"1. {self._clean_text(text)}"
            else:
                # 无序列表
                return f"- {self._clean_text(text)}"

        return self._clean_text(text)

    def _process_table(self, table):
        """
        处理表格，转换为Markdown格式

        Args:
            table: docx.Table对象
        """
        if not table.rows:
            return

        # 获取表格数据
        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = self._clean_text(cell.text)
                row_data.append(cell_text)
            rows_data.append(row_data)

        # 检查是否有数据
        if not rows_data:
            return

        # 第一行作为表头
        header = rows_data[0]
        self.markdown_lines.append("| " + " | ".join(header) + " |")
        self.markdown_lines.append("| " + " | ".join(["---"] * len(header)) + " |")

        # 数据行
        for row in rows_data[1:]:
            self.markdown_lines.append("| " + " | ".join(row) + " |")

        # 表格后添加空行
        self.markdown_lines.append("")

    def convert_paragraphs(self):
        """转换所有段落和表格"""
        # 如果已经从altChunk提取了内容，跳过标准转换
        if self.has_alt_chunk and self.markdown_lines:
            return

        for block in self.iter_block_items():
            if isinstance(block, Paragraph):
                self._process_paragraph(block)
            elif isinstance(block, Table):
                self._process_table(block)

        # 移除连续的空行
        cleaned_lines = []
        prev_empty = False
        for line in self.markdown_lines:
            if line.strip():
                cleaned_lines.append(line)
                prev_empty = False
            elif not prev_empty:
                cleaned_lines.append(line)
                prev_empty = True

        self.markdown_lines = cleaned_lines

    def generate_markdown(self):
        """生成完整的Markdown内容"""
        self.convert_paragraphs()
        return '\n'.join(self.markdown_lines)

    def save(self, output_path=None):
        """
        保存Markdown文件

        Args:
            output_path: 输出文件路径，如果为None则使用输入文件名.md

        Returns:
            保存的文件路径
        """
        if output_path is None:
            output_path = self.docx_path.with_suffix('.md')
        else:
            output_path = Path(output_path)

        # 确保输出目录存在
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # 写入文件
        content = self.generate_markdown()
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

        print(f"成功保存Markdown文件: {output_path}")
        return str(output_path)


def validate_input(docx_path):
    """
    验证输入文件

    Args:
        docx_path: 输入文件路径

    Returns:
        验证通过返回True，否则抛出异常
    """
    path = Path(docx_path)

    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {docx_path}")

    if not path.suffix.lower() in ['.docx', '.doc']:
        raise ValueError(f"不支持的文件格式，请使用.docx文件: {docx_path}")

    return True


def argument_parser():
    """创建命令行参数解析器"""
    parser = argparse.ArgumentParser(
        description='将Word文档(.docx)转换为Markdown格式(.md)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  %(prog)s input.docx                    # 转换为同名.md文件
  %(prog)s input.docx output.md          # 指定输出文件名
  %(prog)s input.docx -o output.md       # 使用-o选项指定输出
        """
    )

    parser.add_argument(
        'input',
        help='输入的Word文档路径(.docx)'
    )

    parser.add_argument(
        'output',
        nargs='?',
        help='输出的Markdown文件路径(可选，默认为同名.md文件)'
    )

    parser.add_argument(
        '-o', '--output',
        dest='output_opt',
        help='指定输出文件路径'
    )

    return parser


def main():
    """主入口函数"""
    # 解析命令行参数
    args = argument_parser().parse_args()

    try:
        # 验证输入文件
        validate_input(args.input)

        # 获取输出路径（优先使用-o选项，其次是第二个位置参数）
        output_path = args.output_opt or args.output

        # 创建转换器并执行转换
        converter = WordToMarkdownConverter(args.input)
        converter.load_document()
        saved_path = converter.save(output_path)

        print(f"\n转换完成!")
        print(f"输入: {args.input}")
        print(f"输出: {saved_path}")

    except FileNotFoundError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"发生未预期的错误: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
